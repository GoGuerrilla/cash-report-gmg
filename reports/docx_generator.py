"""
C.A.S.H. Report by GMG — Word Document Generator
Produces a clean, professional .docx report structured around the
4 C.A.S.H. pillars: Content · Audience · Sales · Hold (Retention).
"""
from datetime import datetime
from typing import Dict, Any, List, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import ClientConfig

# ── Palette ────────────────────────────────────────────────────
NAVY   = "1B2A4A"
GOLD   = "C9A84C"
WHITE  = "FFFFFF"
LGRAY  = "F5F7FA"
DGRAY  = "2C3E50"
GREEN  = "27AE60"
AMBER  = "F39C12"
RED    = "E74C3C"
MGRAY  = "8B9BAB"


def _rgb(hex6: str) -> RGBColor:
    return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))


def _shade_cell(cell, hex_color: str):
    """Fill a table cell with a background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def _cell_para(cell, text: str, bold=False, color=DGRAY, size=9,
               align=WD_ALIGN_PARAGRAPH.LEFT, font="Calibri"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    return para


def _score_color(score: int) -> str:
    if score >= 65: return GREEN
    if score >= 35: return AMBER
    return RED


def _grade(score: int) -> str:
    if score >= 80: return "A"
    if score >= 65: return "B"
    if score >= 50: return "C"
    if score >= 35: return "D"
    return "F"


class DocxReportGenerator:
    def __init__(self, config: ClientConfig, audit_data: Dict[str, Any]):
        self.config = config
        self.data   = audit_data
        self.ai     = audit_data.get("ai_insights", {})
        self.date_str = datetime.now().strftime("%B %d, %Y")
        self.cash   = self._compute_cash_scores()

    # ── Public entry point ─────────────────────────────────────

    def generate(self, output_path: str):
        doc = Document()
        self._set_margins(doc)
        self._set_default_style(doc)

        self._build_title_page(doc)
        self._page_break(doc)
        self._build_intro_page(doc)
        self._page_break(doc)
        self._build_cash_scorecard(doc)
        self._page_break(doc)
        self._build_executive_summary(doc)
        self._page_break(doc)
        self._build_section_c(doc)
        self._page_break(doc)
        self._build_section_a(doc)
        self._page_break(doc)
        self._build_section_s(doc)
        self._page_break(doc)
        self._build_section_h(doc)
        self._page_break(doc)
        self._build_section_geo(doc)
        self._page_break(doc)
        self._build_section_competitive(doc)
        self._page_break(doc)
        self._build_action_plan(doc)
        self._page_break(doc)
        self._build_cta_section(doc)
        self._page_break(doc)
        self._build_appendix(doc)

        doc.save(output_path)

    # ── Document setup ─────────────────────────────────────────

    def _set_margins(self, doc: Document):
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

    def _set_default_style(self, doc: Document):
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.font.color.rgb = _rgb(DGRAY)

    def _page_break(self, doc: Document):
        doc.add_page_break()

    # ── Title page ─────────────────────────────────────────────

    def _build_title_page(self, doc: Document):
        # Report name banner
        banner = doc.add_table(rows=1, cols=1)
        banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        banner.style = "Table Grid"
        bc = banner.rows[0].cells[0]
        _shade_cell(bc, NAVY)
        bc.width = Inches(6.5)
        p = bc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run("C.A.S.H. REPORT")
        r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(28)
        r.font.color.rgb = _rgb(GOLD)
        p2 = bc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(18)
        r2 = p2.add_run("Content  ·  Audience  ·  Sales  ·  Hold (Retention)")
        r2.font.name = "Calibri"; r2.font.size = Pt(11)
        r2.font.color.rgb = _rgb(MGRAY)

        doc.add_paragraph()

        # Client name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self.config.client_name.upper())
        r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(22)
        r.font.color.rgb = _rgb(NAVY)

        # Industry + date
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"{self.config.client_industry}  ·  {self.date_str}")
        r2.font.name = "Calibri"; r2.font.size = Pt(11)
        r2.font.color.rgb = _rgb(MGRAY)

        doc.add_paragraph()

        # Overall grade badge
        overall = self.cash.get("overall", self.ai.get("overall_score", 50))
        grade   = self.ai.get("overall_grade", _grade(overall))
        gc      = _score_color(overall)
        gt = doc.add_table(rows=1, cols=3)
        gt.alignment = WD_TABLE_ALIGNMENT.CENTER
        gt.style = "Table Grid"
        widths = [Inches(2.0), Inches(1.2), Inches(2.0)]
        for i, w in enumerate(widths):
            gt.rows[0].cells[i].width = w
        _shade_cell(gt.rows[0].cells[0], WHITE)
        _shade_cell(gt.rows[0].cells[1], gc)
        _shade_cell(gt.rows[0].cells[2], WHITE)
        gt.rows[0].cells[0].paragraphs[0].text = ""
        _cell_para(gt.rows[0].cells[1], f"{grade}\n{overall}/100",
                   bold=True, color=WHITE, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
        gt.rows[0].cells[2].paragraphs[0].text = ""

        doc.add_paragraph()

        # ICP target
        if self.config.stated_target_market:
            pi = doc.add_paragraph()
            pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ri = pi.add_run(f"Target Market: {self.config.stated_target_market}")
            ri.font.name = "Calibri"; ri.font.size = Pt(10)
            ri.font.color.rgb = _rgb(MGRAY)

        # Prepared by
        doc.add_paragraph()
        self._hairline(doc)
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = pf.add_run(f"Prepared by {self.config.agency_name}")
        rf.font.name = "Calibri"; rf.font.size = Pt(9)
        rf.font.color.rgb = _rgb(MGRAY)

    # ── GMG Introduction page ──────────────────────────────────

    def _build_intro_page(self, doc: Document):
        """
        Page 2 — agency introduction before the scorecard.
        Full navy-and-gold treatment: headline, subheadline, what/why/how,
        C.A.S.H. pillars, and a closing credibility line.
        """

        # ── Top full-width navy banner ─────────────────────────
        banner = doc.add_table(rows=1, cols=1)
        banner.style = "Table Grid"
        banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        bc = banner.rows[0].cells[0]
        _shade_cell(bc, NAVY)
        bc.width = Inches(6.5)

        # Agency label
        p_agency = bc.paragraphs[0]
        p_agency.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_agency.paragraph_format.space_before = Pt(14)
        p_agency.paragraph_format.space_after  = Pt(2)
        r_agency = p_agency.add_run("GUERRILLA MARKETING GROUP  ·  goguerrilla.xyz")
        r_agency.font.name = "Calibri"
        r_agency.font.size = Pt(8)
        r_agency.font.color.rgb = _rgb(MGRAY)
        r_agency.bold = False

        # Headline
        p_hl = bc.add_paragraph()
        p_hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_hl.paragraph_format.space_before = Pt(6)
        p_hl.paragraph_format.space_after  = Pt(4)
        r_hl = p_hl.add_run(
            "Know Exactly Where Your Marketing Is Winning and Losing"
        )
        r_hl.bold = True
        r_hl.font.name = "Calibri"
        r_hl.font.size = Pt(20)
        r_hl.font.color.rgb = _rgb(GOLD)

        # Sub-headline
        p_sub = bc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after  = Pt(16)
        r_sub = p_sub.add_run(
            "An AI-powered audit that scores your entire online presence, reveals your biggest "
            "opportunities, exposes what's wasting your budget, and gives you a clear "
            "90-day plan to grow."
        )
        r_sub.font.name = "Calibri"
        r_sub.font.size = Pt(11)
        r_sub.font.color.rgb = _rgb(WHITE)
        r_sub.bold = False

        doc.add_paragraph()

        # ── What is the C.A.S.H. Report — body paragraph ──────
        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_before = Pt(4)
        p_body.paragraph_format.space_after  = Pt(10)
        r_body = p_body.add_run(
            "The C.A.S.H. Report is a comprehensive marketing audit built by Guerrilla Marketing "
            "Group to give business owners and founders a clear, honest picture of how their "
            "marketing is actually performing — not how they hope it's performing. Most businesses "
            "spend time and money on marketing without knowing which channels are working, "
            "which are wasting budget, and where their biggest growth opportunities are being missed. "
            "This report changes that."
        )
        r_body.font.name = "Calibri"
        r_body.font.size = Pt(10.5)
        r_body.font.color.rgb = _rgb(DGRAY)

        p_body2 = doc.add_paragraph()
        p_body2.paragraph_format.space_before = Pt(0)
        p_body2.paragraph_format.space_after  = Pt(10)
        r_body2 = p_body2.add_run(
            "Every audit runs live checks against your website, social profiles, search "
            "visibility, Google Business Profile, and content output — then cross-references "
            "the findings against your stated target market and ideal client profile. "
            "The result is a scored, prioritised action plan that tells you exactly what to fix "
            "first, what to stop doing, and what to double down on to grow faster."
        )
        r_body2.font.name = "Calibri"
        r_body2.font.size = Pt(10.5)
        r_body2.font.color.rgb = _rgb(DGRAY)

        # ── C.A.S.H. pillars — 4-column table ─────────────────
        doc.add_paragraph()
        pillars_title = doc.add_paragraph()
        pillars_title.paragraph_format.space_before = Pt(6)
        pillars_title.paragraph_format.space_after  = Pt(8)
        r_pt = pillars_title.add_run("THE FOUR C.A.S.H. PILLARS")
        r_pt.bold = True
        r_pt.font.name = "Calibri"
        r_pt.font.size = Pt(9)
        r_pt.font.color.rgb = _rgb(NAVY)

        pillars = [
            ("C", "CONTENT",
             "Is your content attracting the right audience? "
             "We score SEO, freshness, website quality, "
             "and how well your messaging matches your ICP."),
            ("A", "AUDIENCE",
             "Are you reaching the right people on the right channels? "
             "We assess ICP alignment, brand consistency, "
             "and social platform fit for your industry."),
            ("S", "SALES",
             "Does your online presence convert interest into leads? "
             "We evaluate your lead capture, funnel structure, "
             "CTAs, and conversion readiness."),
            ("H", "HOLD",
             "Are you retaining attention and building loyalty? "
             "We audit your nurture systems, email marketing, "
             "referral programmes, and trust signals."),
        ]

        pt = doc.add_table(rows=2, cols=4)
        pt.style = "Table Grid"
        pt.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_w = Inches(1.625)

        # Header row — letter + pillar name
        for ci, (letter, name, _) in enumerate(pillars):
            cell = pt.rows[0].cells[ci]
            cell.width = col_w
            _shade_cell(cell, NAVY)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_before = Pt(8)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
            rl = cell.paragraphs[0].add_run(letter)
            rl.bold = True; rl.font.name = "Calibri"; rl.font.size = Pt(22)
            rl.font.color.rgb = _rgb(GOLD)
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(8)
            rn = p2.add_run(name)
            rn.bold = True; rn.font.name = "Calibri"; rn.font.size = Pt(8)
            rn.font.color.rgb = _rgb(WHITE)

        # Body row — description
        for ci, (_, _, desc) in enumerate(pillars):
            cell = pt.rows[1].cells[ci]
            cell.width = col_w
            _shade_cell(cell, LGRAY)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].paragraph_format.space_before = Pt(6)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(6)
            cell.paragraphs[0].paragraph_format.left_indent  = Pt(4)
            cell.paragraphs[0].paragraph_format.right_indent = Pt(4)
            rd = cell.paragraphs[0].add_run(desc)
            rd.font.name = "Calibri"; rd.font.size = Pt(8.5)
            rd.font.color.rgb = _rgb(DGRAY)

        doc.add_paragraph()

        # ── What this report gives you — two-column checklist ──
        delivers_title = doc.add_paragraph()
        delivers_title.paragraph_format.space_before = Pt(4)
        delivers_title.paragraph_format.space_after  = Pt(6)
        r_dt = delivers_title.add_run("WHAT THIS REPORT GIVES YOU")
        r_dt.bold = True
        r_dt.font.name = "Calibri"
        r_dt.font.size = Pt(9)
        r_dt.font.color.rgb = _rgb(NAVY)

        deliverables = [
            ("A scored audit across 8 marketing dimensions",
             "Side-by-side competitor comparison"),
            ("Your overall C.A.S.H. grade (A–F) with component breakdown",
             "GEO score — how visible you are to AI tools like ChatGPT"),
            ("The single biggest waste in your current marketing spend",
             "ICP alignment verdict — are you reaching the right people"),
            ("Your highest-ROI opportunity, ranked and explained",
             "A clear, week-by-week 90-day action plan"),
        ]

        dt = doc.add_table(rows=len(deliverables), cols=2)
        dt.style = "Table Grid"
        dt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, (left, right) in enumerate(deliverables):
            row_bg = LGRAY if ri % 2 == 0 else WHITE
            for ci, text in enumerate([left, right]):
                cell = dt.rows[ri].cells[ci]
                cell.width = Inches(3.25)
                _shade_cell(cell, row_bg)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                p.paragraph_format.left_indent  = Pt(4)
                tick = cell.paragraphs[0].add_run("✓  ")
                tick.bold = True
                tick.font.name = "Calibri"; tick.font.size = Pt(9)
                tick.font.color.rgb = _rgb(GOLD)
                rd = cell.paragraphs[0].add_run(text)
                rd.font.name = "Calibri"; rd.font.size = Pt(9)
                rd.font.color.rgb = _rgb(DGRAY)

        doc.add_paragraph()

        # ── Closing credibility strip ──────────────────────────
        cred = doc.add_table(rows=1, cols=1)
        cred.style = "Table Grid"
        cred.alignment = WD_TABLE_ALIGNMENT.CENTER
        cc = cred.rows[0].cells[0]
        _shade_cell(cc, NAVY)
        cc.width = Inches(6.5)
        pc = cc.paragraphs[0]
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(10)
        pc.paragraph_format.space_after  = Pt(10)
        rc = pc.add_run(
            "Built by GMG — the marketing partner for businesses that want results, not reports that sit in a drawer.  "
            "Every finding in this document is based on live data pulled at the time of audit."
        )
        rc.font.name = "Calibri"
        rc.font.size = Pt(9)
        rc.font.color.rgb = _rgb(MGRAY)

    # ── CASH Scorecard ─────────────────────────────────────────

    def _build_cash_scorecard(self, doc: Document):
        self._section_header(doc, "C.A.S.H. SCORE OVERVIEW")

        overall = self.cash.get("overall", self.ai.get("overall_score", 50))

        # 4-box CASH scores
        labels = [
            ("C", "Content",          self.cash.get("C", 50)),
            ("A", "Audience",         self.cash.get("A", 50)),
            ("S", "Sales",            self.cash.get("S", 50)),
            ("H", "Hold / Retention", self.cash.get("H", 50)),
        ]

        t = doc.add_table(rows=2, cols=4)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        for col_idx, (letter, label, score) in enumerate(labels):
            hdr = t.rows[0].cells[col_idx]
            val = t.rows[1].cells[col_idx]
            hdr.width = val.width = Inches(1.5)
            sc = _score_color(score)
            _shade_cell(hdr, NAVY)
            _shade_cell(val, sc)
            _cell_para(hdr, f"{letter} — {label}", bold=True, color=WHITE,
                       size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_para(val, f"{score}/100  ({_grade(score)})", bold=True, color=WHITE,
                       size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()

        # Overall score row
        ot = doc.add_table(rows=1, cols=2)
        ot.alignment = WD_TABLE_ALIGNMENT.CENTER
        ot.style = "Table Grid"
        lc = ot.rows[0].cells[0]
        vc = ot.rows[0].cells[1]
        lc.width = Inches(3.0); vc.width = Inches(3.0)
        _shade_cell(lc, NAVY); _shade_cell(vc, _score_color(overall))
        _cell_para(lc, "OVERALL C.A.S.H. SCORE", bold=True, color=WHITE,
                   size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(vc, f"{overall}/100  ({self.ai.get('overall_grade', _grade(overall))})",
                   bold=True, color=WHITE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()

        # Component detail table
        self._subsection(doc, "Component Breakdown")
        web = self.data.get("website", {}).get("scores", {})
        rows = [
            ("ICP / Audience Alignment",   self.data.get("icp", {}).get("score", 50),     "A"),
            ("Brand Consistency",          self.data.get("brand", {}).get("score", 50),   "A"),
            ("Content Freshness",          self.data.get("freshness", {}).get("score", 50), "C"),
            ("SEO Health",                 self.data.get("seo", {}).get("score", 50),       "C"),
            ("Lead Capture (Sales)",       self.cash.get("S", 50),                          "S"),
            ("Retention / Hold Systems",   self.cash.get("H", 50),                          "H"),
            ("Website Technical",          web.get("technical", 50),                        "—"),
            ("Website Conversion",         web.get("conversion", 50),                       "S"),
        ]
        self._detail_table(doc, ["COMPONENT", "SCORE", "PILLAR", "STATUS"], [
            (label, f"{score}/100", pillar,
             "Good" if score >= 70 else "Needs Work" if score >= 50 else "Critical")
            for label, score, pillar in rows
        ], col_widths=[Inches(2.8), Inches(0.9), Inches(0.8), Inches(1.5)])

    # ── Executive Summary ──────────────────────────────────────

    def _build_executive_summary(self, doc: Document):
        self._section_header(doc, "EXECUTIVE SUMMARY")

        summary = self.ai.get("executive_summary", "")
        if summary:
            p = doc.add_paragraph()
            r = p.add_run(summary)
            r.font.name = "Calibri"; r.font.size = Pt(11)
            r.font.color.rgb = _rgb(DGRAY)
            doc.add_paragraph()

        # Opportunity / Waste side by side
        opp   = self.ai.get("biggest_opportunity", "")
        waste = self.ai.get("biggest_waste", "")
        if opp or waste:
            t = doc.add_table(rows=1, cols=2)
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            lc = t.rows[0].cells[0]; rc = t.rows[0].cells[1]
            lc.width = rc.width = Inches(3.0)
            _shade_cell(lc, "E8F5E9"); _shade_cell(rc, "FFF3E0")
            p1 = lc.paragraphs[0]
            p1.paragraph_format.space_before = Pt(4)
            p1.paragraph_format.space_after  = Pt(4)
            r1 = p1.add_run("BIGGEST OPPORTUNITY\n")
            r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(9)
            r1.font.color.rgb = _rgb(GREEN)
            r1b = p1.add_run(opp or "—")
            r1b.font.name = "Calibri"; r1b.font.size = Pt(9)
            r1b.font.color.rgb = _rgb(DGRAY)
            p2 = rc.paragraphs[0]
            p2.paragraph_format.space_before = Pt(4)
            p2.paragraph_format.space_after  = Pt(4)
            r2 = p2.add_run("BIGGEST WASTE\n")
            r2.bold = True; r2.font.name = "Calibri"; r2.font.size = Pt(9)
            r2.font.color.rgb = _rgb(RED)
            r2b = p2.add_run(waste or "—")
            r2b.font.name = "Calibri"; r2b.font.size = Pt(9)
            r2b.font.color.rgb = _rgb(DGRAY)
            doc.add_paragraph()

        # Top 3 priorities
        priorities = self.ai.get("top_3_priorities", [])
        if priorities:
            self._subsection(doc, "Top 3 Priorities")
            self._detail_table(doc, ["#", "ACTION", "IMPACT", "TIMELINE"], [
                (str(p.get("priority", i+1)),
                 p.get("action", ""),
                 p.get("impact", ""),
                 p.get("timeline", ""))
                for i, p in enumerate(priorities[:3])
            ], col_widths=[Inches(0.3), Inches(2.2), Inches(2.5), Inches(1.0)])

        # ICP alignment verdict
        verdict = self.ai.get("icp_alignment_verdict", "")
        if not verdict:
            verdict = self.data.get("icp", {}).get("icp_verdict", "")
        if verdict:
            doc.add_paragraph()
            self._subsection(doc, "ICP Alignment Verdict")
            self._callout(doc, verdict, AMBER)

    # ── Section C: Content ─────────────────────────────────────

    def _build_section_c(self, doc: Document):
        self._cash_section_banner(doc, "C", "CONTENT", self.cash.get("C", 50),
            "How fresh, consistent, and strategically distributed is the content?")

        # Industry benchmark context note
        industry = getattr(self.config, "industry_category", "") or self.config.client_industry
        if industry and industry != "Other":
            p = doc.add_paragraph()
            r = p.add_run(f"Benchmarks calibrated for: {industry}")
            r.font.name = "Calibri"; r.font.size = Pt(9); r.italic = True
            r.font.color.rgb = _rgb(MGRAY)
            p.paragraph_format.space_after = Pt(4)

        # Freshness by channel
        freshness = self.data.get("freshness", {})
        channels  = freshness.get("channels", {})
        if channels:
            self._subsection(doc, "Content Freshness by Channel")
            status_map = {
                "fresh":            "✅ Fresh",
                "recent":           "✅ Recent",
                "stale":            "⚠️ Stale",
                "dead":             "🔴 Inactive",
                "unknown":          "— Unknown",
                "unknown_inactive": "— Unverified",
                "api_blocked":      "— Requires intake or API",
            }
            self._detail_table(doc, ["CHANNEL", "STATUS", "POSTS/WEEK", "DAYS SINCE POST"], [
                (p,
                 status_map.get(d.get("status", "unknown"), "—"),
                 str(d.get("posts_per_week") or "?"),
                 str(d.get("days_since_last_post") or "?"))
                for p, d in channels.items()
            ], col_widths=[Inches(1.3), Inches(1.4), Inches(1.4), Inches(1.9)])

        # YouTube channel metrics (live via YouTube Data API v3)
        content = self.data.get("content", {})
        yt = content.get("youtube_metrics")
        if yt and yt.get("data_source") == "youtube_api_v3":
            self._subsection(doc, "YouTube Channel Metrics  (Live — YouTube Data API v3)")
            rows = [
                ("Subscribers",              f"{yt.get('subscriber_count', 0):,}"),
                ("Total Videos Published",   f"{yt.get('total_video_count', 0):,}"),
                ("Total Channel Views",      f"{yt.get('total_view_count', 0):,}"),
                ("Videos Uploaded (30 Days)", str(yt.get("videos_last_30_days", 0))),
                ("Upload Rate",              f"{yt.get('posts_per_week', 0)} videos/week"),
                ("Avg Views per Video",      f"{yt.get('avg_views_per_video', 0):,}"),
                ("Days Since Last Upload",   str(yt.get("days_since_last_post", "—"))),
                ("Most Viewed Video",        yt.get("most_viewed_video_title") or "—"),
            ]
            self._detail_table(doc, ["METRIC", "VALUE"], rows,
                               col_widths=[Inches(2.4), Inches(3.6)])
            desc = (yt.get("description") or "").strip()
            if desc:
                p = doc.add_paragraph()
                r = p.add_run(f"Channel Description: {desc[:300]}")
                r.font.name = "Calibri"
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(MGRAY)

        # SEO checks
        seo = self.data.get("seo", {})
        if seo:
            self._subsection(doc, "SEO Health")
            checks = [
                ("robots.txt",        seo.get("robots_txt", {}).get("exists", False)),
                ("XML Sitemap",       seo.get("sitemap", {}).get("found", False)),
                ("Canonical Tags",    seo.get("canonical", {}).get("present", False)),
                ("Open Graph Tags",   seo.get("open_graph", {}).get("present", False)),
                ("OG Image Tag",      seo.get("open_graph", {}).get("has_og_image", False)),
            ]
            self._detail_table(doc, ["SEO CHECK", "STATUS"], [
                (label, "✅ Pass" if passed else "❌ Fail")
                for label, passed in checks
            ], col_widths=[Inches(3.5), Inches(2.5)])

        # Website traffic (Google Analytics Data API v4)
        analytics = self.data.get("analytics", {})
        if analytics.get("data_source") == "google_analytics_data_api_v4":
            self._subsection(doc, "Website Traffic  (Live — Google Analytics 4)")
            trend_pct = analytics.get("traffic_trend_pct")
            trend_str = analytics.get("traffic_trend_label", "—")
            br        = analytics.get("bounce_rate_pct")
            self._detail_table(doc, ["METRIC", "VALUE"], [
                ("Monthly Visitors (30 Days)",    f"{analytics.get('monthly_visitors', 0):,}"),
                ("Traffic Trend (vs Prior 30d)",  trend_str),
                ("Bounce Rate",                   f"{br}%" if br is not None else "—"),
                ("Avg Session Duration",          analytics.get("avg_session_duration", "—")),
            ], col_widths=[Inches(2.4), Inches(3.6)])

            top_src = analytics.get("top_traffic_sources", [])
            if top_src:
                self._subsection(doc, "Top Traffic Sources")
                self._detail_table(doc, ["CHANNEL", "SESSIONS"], [
                    (s.get("channel", "—"), f"{s.get('sessions', 0):,}")
                    for s in top_src
                ], col_widths=[Inches(3.5), Inches(2.5)])

            top_pages = analytics.get("top_landing_pages", [])
            if top_pages:
                self._subsection(doc, "Top Landing Pages")
                self._detail_table(doc, ["PAGE", "SESSIONS", "BOUNCE RATE"], [
                    (p.get("page", "—")[:60],
                     f"{p.get('sessions', 0):,}",
                     f"{p.get('bounce_rate', 0):.0f}%")
                    for p in top_pages
                ], col_widths=[Inches(3.2), Inches(1.2), Inches(1.6)])
        elif analytics.get("note"):
            self._subsection(doc, "Website Traffic")
            self._callout(doc, f"⚠️  {analytics['note']}", MGRAY)

        # Content gaps & quick wins
        content = self.data.get("content", {})
        gaps    = content.get("content_gaps", [])
        wins    = content.get("quick_wins", [])

        if gaps:
            self._subsection(doc, "Content Gaps")
            self._detail_table(doc, ["GAP", "IMPACT", "FIX"], [
                (g.get("gap", ""), g.get("impact", ""), g.get("fix", ""))
                for g in gaps
            ], col_widths=[Inches(1.6), Inches(2.2), Inches(2.2)])

        if wins:
            self._subsection(doc, "Quick Wins")
            self._detail_table(doc, ["WIN", "EFFORT", "IMPACT", "WHEN"], [
                (w.get("win", ""), w.get("effort", ""), w.get("impact", ""), w.get("timeline", ""))
                for w in wins
            ], col_widths=[Inches(2.0), Inches(1.2), Inches(1.8), Inches(1.0)])

        # API-blocked platforms note
        scraping_note = freshness.get("scraping_note")
        if scraping_note:
            self._subsection(doc, "Data Availability Note")
            self._callout(doc, scraping_note, MGRAY)

        # Issues & strengths
        self._issues_strengths(doc,
            freshness.get("issues", []) + seo.get("issues", []),
            freshness.get("strengths", []) + seo.get("strengths", []))

        # Content strategy recommendation
        cs = self.ai.get("content_strategy", "")
        if cs:
            self._subsection(doc, "Content Strategy Recommendation")
            self._callout(doc, cs, NAVY)

    # ── Section A: Audience ────────────────────────────────────

    def _build_section_a(self, doc: Document):
        self._cash_section_banner(doc, "A", "AUDIENCE", self.cash.get("A", 50),
            "Are you reaching the right people on the right platforms?")

        icp   = self.data.get("icp", {})
        brand = self.data.get("brand", {})

        # ICP verdict callout
        verdict = icp.get("icp_verdict", "")
        if verdict:
            self._subsection(doc, "ICP Alignment Verdict")
            self._callout(doc, verdict, _score_color(icp.get("score", 50)))

        # Platform fit table
        pf = brand.get("platform_fit", {})
        ps = pf.get("platform_scores", {})
        if ps:
            self._subsection(doc, "Platform Fit for Target Market")
            high = pf.get("high_fit", [])
            med  = pf.get("medium_fit", [])
            self._detail_table(doc, ["PLATFORM", "FIT SCORE", "RECOMMENDATION"], [
                (p,
                 f"{sc}/100",
                 "✅ Prioritize" if p in high else ("⚠️ Use selectively" if p in med else "🔴 Deprioritize"))
                for p, sc in sorted(ps.items(), key=lambda x: -x[1])
            ], col_widths=[Inches(1.5), Inches(1.2), Inches(3.3)])

        # Issues & strengths
        self._issues_strengths(doc,
            icp.get("issues", []) + brand.get("issues", []),
            icp.get("strengths", []) + brand.get("strengths", []))

        # ICP recommendations
        recs = icp.get("recommendations", [])
        if recs:
            self._subsection(doc, "Audience & ICP Recommendations")
            self._detail_table(doc, ["PRIORITY", "ACTION", "DETAIL", "TIMELINE"], [
                (r.get("priority", ""), r.get("action", ""), r.get("detail", ""), r.get("timeline", ""))
                for r in recs
            ], col_widths=[Inches(0.8), Inches(1.7), Inches(2.5), Inches(1.0)])

        cr = self.ai.get("channel_recommendation", "")
        if cr:
            self._subsection(doc, "Channel Strategy")
            self._callout(doc, cr, NAVY)

    # ── Section S: Sales ───────────────────────────────────────

    def _build_section_s(self, doc: Document):
        self._cash_section_banner(doc, "S", "SALES", self.cash.get("S", 50),
            "How effectively does the brand convert attention into qualified leads?")

        funnel = self.data.get("funnel", {})
        stages = funnel.get("stages", {})

        # Stage status table
        if stages:
            self._subsection(doc, "Funnel Stage Analysis")
            stage_labels = [
                ("awareness",  "Awareness"),
                ("capture",    "Lead Capture"),
                ("conversion", "Conversion"),
            ]
            rows = []
            for key, label in stage_labels:
                stage = stages.get(key, {})
                n_crit = len([i for i in stage.get("issues", []) if "🔴" in i])
                n_str  = len(stage.get("strengths", []))
                status = "🔴 Critical" if n_crit >= 2 else ("⚠️ Needs Work" if n_crit == 1 else "✅ OK")
                rows.append((label, str(n_crit), str(n_str), status))
            self._detail_table(doc, ["STAGE", "CRITICAL ISSUES", "STRENGTHS", "STATUS"],
                               rows, col_widths=[Inches(1.8), Inches(1.4), Inches(1.4), Inches(1.4)])

        # Issues & strengths (capture + conversion only)
        capture    = stages.get("capture", {})
        conversion = stages.get("conversion", {})
        web_issues   = self.data.get("website", {}).get("issues", [])
        web_strengths = self.data.get("website", {}).get("strengths", [])
        self._issues_strengths(doc,
            capture.get("issues", []) + conversion.get("issues", []) + web_issues,
            capture.get("strengths", []) + conversion.get("strengths", []) + web_strengths)

        # Funnel recommendations
        recs = funnel.get("recommendations", [])
        if recs:
            self._subsection(doc, "Sales Funnel Recommendations")
            self._detail_table(doc, ["PRIORITY", "ACTION", "EXAMPLE / DETAIL", "TIMELINE"], [
                (r.get("priority", ""), r.get("action", ""),
                 r.get("example", r.get("detail", "")), r.get("timeline", ""))
                for r in recs
            ], col_widths=[Inches(0.7), Inches(1.8), Inches(2.5), Inches(1.0)])

        br = self.ai.get("budget_recommendation", "")
        if br:
            self._subsection(doc, "Budget Recommendation")
            self._callout(doc, br, GOLD)

    # ── Section H: Hold / Retention ───────────────────────────

    def _build_section_h(self, doc: Document):
        self._cash_section_banner(doc, "H", "HOLD (RETENTION)", self.cash.get("H", 50),
            "Are systems in place to retain clients and generate referrals?")

        funnel = self.data.get("funnel", {})
        stages = funnel.get("stages", {})
        nurture = stages.get("nurture", {})
        trust   = stages.get("trust", {})

        # Retention system status
        self._subsection(doc, "Retention System Status")
        items = [
            ("Email Newsletter",    "✅ Active" if self.config.has_active_newsletter else "🔴 Not found"),
            ("Email List Size",     f"{self.config.email_list_size:,} contacts" if self.config.email_list_size > 0 else "🔴 None reported"),
            ("Referral System",     f"✅ {self.config.referral_system_description or 'Yes'}" if self.config.has_referral_system else "🔴 None detected"),
            ("Current Clients",     f"{self.config.current_client_count} ({self.config.current_client_types or 'types unknown'})" if self.config.current_client_count else "Not reported"),
            ("Blog / Content Hub",  "✅ Present" if self.data.get("content", {}).get("content_gaps") is not None
                                                 and not any("blog" in g.get("gap","").lower()
                                                             for g in self.data.get("content",{}).get("content_gaps",[]))
                                                 else "🔴 Not found"),
        ]
        self._detail_table(doc, ["RETENTION ELEMENT", "STATUS"],
                           items, col_widths=[Inches(2.5), Inches(3.5)])

        # Issues & strengths
        self._issues_strengths(doc,
            nurture.get("issues", []) + trust.get("issues", []),
            nurture.get("strengths", []) + trust.get("strengths", []))

        # Hold recommendations
        self._subsection(doc, "Retention Recommendations")
        hold_recs = [
            ("HIGH",   "Build a 5-email welcome sequence for new leads",
             "Introduce the agency, share a case study, and book a discovery call.",
             "2 weeks"),
            ("HIGH",   "Launch a client referral program",
             "Offer 1 free month of service for every referred client who signs on.",
             "1 month"),
            ("MEDIUM", "Start a biweekly email newsletter for financial advisors, CPAs, and attorneys",
             "Content: compliance-safe content tips, case studies, industry news tailored to professional services.",
             "2-4 weeks"),
            ("MEDIUM", "Collect testimonials from every current client",
             "Ask for a 2-3 sentence quote + permission to use on website and LinkedIn.",
             "2 weeks"),
        ]
        self._detail_table(doc, ["PRIORITY", "ACTION", "DETAIL", "TIMELINE"],
                           hold_recs, col_widths=[Inches(0.7), Inches(2.0), Inches(2.4), Inches(0.9)])

    # ── GEO: Generative Engine Optimization ───────────────────

    def _build_section_geo(self, doc: Document):
        geo = self.data.get("geo", {})
        score = geo.get("score", 50)
        self._cash_section_banner(doc, "GEO", "GENERATIVE ENGINE OPTIMIZATION", score,
            "How visible is this brand in AI-generated answers (ChatGPT, Google AI Overviews, Perplexity)?")

        # Component breakdown table
        components = geo.get("components", {})
        if components:
            self._subsection(doc, "GEO Component Scores")
            weights = {
                "Schema Markup":     "25%",
                "FAQ / Q&A Content": "20%",
                "E-E-A-T Signals":   "20%",
                "Brand Authority":   "20%",
                "AI Citation Score": "15%",
            }
            self._detail_table(doc, ["COMPONENT", "WEIGHT", "SCORE", "GRADE"], [
                (name,
                 weights.get(name, "—"),
                 f"{sc}/100",
                 _grade(sc))
                for name, sc in components.items()
            ], col_widths=[Inches(2.4), Inches(0.8), Inches(0.9), Inches(0.9)])

        # Google Business Profile
        gbp = self.data.get("gbp", {})
        if gbp.get("found"):
            self._subsection(doc, "Google Business Profile  (Live — Google Places API)")
            rating_str = f"{gbp['rating']}/5" if gbp.get("rating") else "No rating"
            rows = [
                ("Business Name",      gbp.get("business_name", "—")),
                ("Address",            gbp.get("address", "—")),
                ("Phone",              gbp.get("phone") or "—"),
                ("Rating",             rating_str),
                ("Total Reviews",      str(gbp.get("review_count", 0))),
                ("Photos on Profile",  str(gbp.get("photo_count", 0))),
                ("Hours Listed",       "✅ Yes" if gbp.get("hours_listed") else "❌ No"),
                ("Profile Status",     gbp.get("business_status", "—")),
                ("Appears Verified",   "✅ Yes" if gbp.get("is_likely_verified") else "⚠️ Uncertain"),
                ("GBP Posts",          "Cannot verify via Places API"),
                ("Profile Complete",   f"{gbp.get('completeness_pct', 0)}%"),
                ("GBP Score",          f"{gbp.get('score', 50)}/100  ({gbp.get('grade', 'C')})"),
            ]
            self._detail_table(doc, ["FIELD", "VALUE"], rows,
                               col_widths=[Inches(2.0), Inches(4.0)])

            # Score breakdown
            breakdown = gbp.get("score_breakdown", {})
            if breakdown:
                self._subsection(doc, "GBP Score Breakdown")
                self._detail_table(doc, ["COMPONENT", "MAX PTS", "SCORED"], [
                    ("Listing Confirmed", "35", str(breakdown.get("listing_confirmed", 0))),
                    ("NAP on Website",    "25", str(breakdown.get("nap_on_website",    0))),
                    ("Reviews Promoted",  "25", str(breakdown.get("reviews_promoted",  0))),
                    ("Schema Markup",     "15", str(breakdown.get("schema_markup",     0))),
                    ("NAP Consistent",    "5",  str(breakdown.get("nap_consistent",    0))),
                ], col_widths=[Inches(2.2), Inches(1.0), Inches(2.8)])

            self._issues_strengths(doc, gbp.get("issues", []), gbp.get("strengths", []))

        elif gbp.get("note"):
            self._subsection(doc, "Google Business Profile")
            self._callout(doc, f"⚠️  {gbp['note']}", MGRAY)

        # Issues & strengths
        self._issues_strengths(doc, geo.get("issues", []), geo.get("strengths", []))

        # Platform-specific notes
        platform_notes = geo.get("platform_notes", {})
        if platform_notes:
            self._subsection(doc, "AI Platform Visibility Notes")
            for platform, note in platform_notes.items():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(2)
                r_label = p.add_run(f"{platform}:  ")
                r_label.bold = True; r_label.font.name = "Calibri"
                r_label.font.size = Pt(9); r_label.font.color.rgb = _rgb(NAVY)
                r_text = p.add_run(note)
                r_text.font.name = "Calibri"; r_text.font.size = Pt(9)
                r_text.font.color.rgb = _rgb(DGRAY)
            doc.add_paragraph()

        # GEO recommendations table
        recs = geo.get("recommendations", [])
        if recs:
            self._subsection(doc, "GEO Recommendations")
            self._detail_table(doc, ["PRIORITY", "ACTION", "DETAIL", "TIMELINE", "IMPACT"], [
                (r.get("priority", ""),
                 r.get("action", ""),
                 r.get("detail", ""),
                 r.get("timeline", ""),
                 r.get("impact", ""))
                for r in recs
            ], col_widths=[Inches(0.7), Inches(1.4), Inches(2.0), Inches(0.8), Inches(1.6)])

    # ── 90-Day Action Plan ─────────────────────────────────────

    # ── Competitive Positioning Analysis ──────────────────────

    def _build_section_competitive(self, doc: Document):
        """Side-by-side comparison of client vs up to 3 competitor sites."""
        self._section_header(doc, "COMPETITIVE POSITIONING ANALYSIS")

        comp_data = self.data.get("competitor", {})

        # Biggest marketing challenge
        challenge = (
            self.config.biggest_marketing_challenge
            or comp_data.get("biggest_challenge", "")
        )
        if challenge:
            doc.add_paragraph()
            self._subsection(doc, "Biggest Marketing Challenge")
            self._callout(doc, challenge, NAVY)

        if comp_data.get("skipped") or not comp_data.get("competitors"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            r = p.add_run(
                "No competitor URLs were provided. Add up to 3 competitor websites "
                "in the intake questionnaire to unlock this section."
            )
            r.font.name  = "Calibri"
            r.font.size  = Pt(10)
            r.font.color.rgb = _rgb(MGRAY)
            r.italic     = True
            return

        comparison = comp_data.get("comparison", {})
        competitors = comparison.get("competitors", [])
        client      = comparison.get("client", {})
        rows        = comparison.get("rows", [])

        if not competitors:
            return

        # ── Competitor URL reference ────────────────────────────
        doc.add_paragraph()
        self._subsection(doc, "Competitors Analysed")
        for i, comp in enumerate(competitors, 1):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"Competitor {i}: {comp.get('domain', comp.get('url', ''))}")
            r.font.name  = "Calibri"
            r.font.size  = Pt(10)
            r.font.color.rgb = _rgb(DGRAY)
            note = comp.get("note", "")
            if note:
                rn = p.add_run(f"  ({note})")
                rn.font.name  = "Calibri"
                rn.font.size  = Pt(9)
                rn.font.color.rgb = _rgb(MGRAY)
                rn.italic = True

        # ── Side-by-side comparison table ─────────────────────
        doc.add_paragraph()
        self._subsection(doc, "Side-by-Side Score Comparison")

        n_comps = len(competitors)
        n_cols  = 1 + 1 + n_comps   # Metric | Client | Comp1 [| Comp2 | Comp3]

        # Column widths
        metric_w = Inches(1.7)
        client_w = Inches(1.1)
        comp_w   = Inches(1.1)
        table = doc.add_table(rows=1, cols=n_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # ── Header row ─────────────────────────────────────────
        hdr_cells = table.rows[0].cells
        _shade_cell(hdr_cells[0], NAVY)
        _cell_para(hdr_cells[0], "Metric", bold=True, color=WHITE, size=9,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        hdr_cells[0].width = metric_w

        _shade_cell(hdr_cells[1], GOLD)
        client_label = (client.get("label") or self.config.client_name)[:18]
        _cell_para(hdr_cells[1], client_label, bold=True, color=WHITE, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        hdr_cells[1].width = client_w

        for ci, comp in enumerate(competitors):
            cell = hdr_cells[2 + ci]
            _shade_cell(cell, NAVY)
            label = (comp.get("domain", f"Competitor {ci+1}"))[:16]
            _cell_para(cell, label, bold=True, color=WHITE, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            cell.width = comp_w

        # ── Data rows ──────────────────────────────────────────
        metrics_to_show = [
            ("seo_score",           "SEO Score",           True),
            ("performance_score",   "Performance",         True),
            ("technical_score",     "Website Technical",   True),
            ("content_score",       "Website Content",     True),
            ("conversion_score",    "Website Conversion",  True),
            ("social_channel_count","Social Channels",     False),
        ]

        for row_i, (key, label, is_score) in enumerate(metrics_to_show):
            row = table.add_row()
            cells = row.cells

            # Alternating row background
            row_bg = LGRAY if row_i % 2 == 0 else WHITE
            for cell in cells:
                _shade_cell(cell, row_bg)

            # Metric label cell
            cells[0].width = metric_w
            _cell_para(cells[0], label, bold=False, color=DGRAY, size=9)

            # Client value cell
            client_val = client.get(key, 50)
            comp_vals  = [c.get(key, 50) for c in competitors]
            cells[1].width = client_w

            if is_score:
                disp = f"{client_val}/100"
                color = _score_color(client_val) if isinstance(client_val, int) else MGRAY
            else:
                disp  = str(client_val)
                color = GREEN if isinstance(client_val, int) and client_val >= 3 else AMBER
            _cell_para(cells[1], disp, bold=True, color=color, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

            # Competitor value cells
            for ci, comp in enumerate(competitors):
                cell  = cells[2 + ci]
                cell.width = comp_w
                val   = comp.get(key, 50)
                if is_score:
                    disp  = f"{val}/100" if isinstance(val, int) else "—"
                    color = _score_color(val) if isinstance(val, int) else MGRAY
                else:
                    disp  = str(val) if val else "—"
                    color = GREEN if isinstance(val, int) and val >= 3 else AMBER
                _cell_para(cell, disp, bold=False, color=color, size=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER)

        # ── Key Insights ───────────────────────────────────────
        insights = comp_data.get("insights", [])
        if insights:
            doc.add_paragraph()
            self._subsection(doc, "Key Competitive Insights")
            for insight in insights:
                p = doc.add_paragraph(style="List Bullet")
                clean = insight.lstrip("✅🟡🔴 ")
                is_strength = insight.startswith("✅")
                r = p.add_run(clean)
                r.font.name  = "Calibri"
                r.font.size  = Pt(10)
                r.font.color.rgb = _rgb(GREEN if is_strength else DGRAY)

        # ── AI competitive positioning narrative ───────────────
        ai_cp = self.ai.get("competitive_positioning", "")
        if ai_cp:
            doc.add_paragraph()
            self._subsection(doc, "Strategic Positioning Recommendation")
            self._callout(doc, ai_cp, NAVY)

    def _build_action_plan(self, doc: Document):
        self._section_header(doc, "90-DAY ACTION PLAN")

        plan = self.ai.get("90_day_action_plan", [])
        if plan and isinstance(plan[0], dict):
            self._detail_table(doc, ["WEEKS", "ACTION", "EXPECTED OUTCOME"], [
                (item.get("week", ""), item.get("action", ""), item.get("outcome", ""))
                for item in plan
            ], col_widths=[Inches(0.9), Inches(3.0), Inches(2.1)])
        else:
            # Fallback: plain list
            plan_items = self.ai.get("30_day_action_plan", plan)
            for item in plan_items:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(str(item))
                r.font.name = "Calibri"; r.font.size = Pt(10)
                r.font.color.rgb = _rgb(DGRAY)

        cp = self.ai.get("competitive_positioning", "")
        if cp:
            doc.add_paragraph()
            self._subsection(doc, "Competitive Positioning")
            self._callout(doc, cp, NAVY)

    # ── Appendix ───────────────────────────────────────────────

    def _build_cta_section(self, doc: Document):
        """Final call-to-action page."""
        # Section header
        self._section_header(doc, "READY TO IMPROVE YOUR C.A.S.H. SCORE?")

        doc.add_paragraph()

        # Main value statement
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "GMG offers Fractional CMO services for professional services firms — "
            "financial advisors, CPAs, attorneys, law firms, fractional CFOs, "
            "and corporations seeking marketing leadership without the full-time hire."
        )
        r.font.name = "Calibri"
        r.font.size = Pt(13)
        r.font.color.rgb = _rgb(DGRAY)

        doc.add_paragraph()

        # What's included block
        self._subsection(doc, "What's Included in a GMG Fractional CMO Engagement")
        bullets = [
            "Full C.A.S.H. audit with quarterly re-scoring",
            "ICP-specific content strategy and editorial calendar",
            "LinkedIn authority-building and engagement system",
            "Lead magnet creation and email nurture sequence",
            "Compliance-aware messaging for regulated professionals",
            "Monthly KPI reporting and 90-day action plan updates",
        ]
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(b)
            r.font.name = "Calibri"; r.font.size = Pt(11)
            r.font.color.rgb = _rgb(DGRAY)

        doc.add_paragraph()

        # CTA box — shaded table
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        _shade_cell(cell, NAVY)
        cell.width = Inches(5.5)

        # Book a call line
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(14)
        r1 = p1.add_run("Book a FREE 30-Minute Strategy Call")
        r1.font.name = "Calibri"; r1.font.size = Pt(14)
        r1.bold = True
        r1.font.color.rgb = _rgb(WHITE)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("[YOUR BOOKING LINK]")
        r2.font.name = "Calibri"; r2.font.size = Pt(13)
        r2.font.color.rgb = _rgb(GOLD)

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(14)
        r3 = p3.add_run("No obligation. No pitch deck. Just a clear plan for your firm's marketing.")
        r3.font.name = "Calibri"; r3.font.size = Pt(10)
        r3.font.color.rgb = _rgb(MGRAY)

        doc.add_paragraph()

        # Footer note
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = pf.add_run(
            "C.A.S.H. Report prepared by Guerrilla Marketing Group  ·  goguerrilla.xyz  ·  "
            f"{self.date_str}"
        )
        rf.font.name = "Calibri"; rf.font.size = Pt(9)
        rf.font.color.rgb = _rgb(MGRAY)

    def _build_appendix(self, doc: Document):
        self._section_header(doc, "APPENDIX")

        # Channels audited
        self._subsection(doc, "Channels Audited")
        lt = self.data.get("linktree", {})
        channels_found = lt.get("platforms_found", self.config.active_social_channels)
        self._detail_table(doc, ["CHANNEL", "URL / HANDLE"], [
            (p, urls[0] if isinstance(urls, list) else str(urls))
            for p, urls in lt.get("classified_links", {}).items()
        ] if lt.get("classified_links") else [
            (c, "Configured") for c in channels_found
        ], col_widths=[Inches(1.5), Inches(4.5)])

        doc.add_paragraph()
        self._subsection(doc, "Methodology & Data Confidence")
        notes = [
            "Scores are composite metrics derived from publicly accessible channel data.",
            "Channels that block public scraping (Instagram, Facebook) are scored "
            "at 50 (neutral). YouTube is scored from live YouTube Data API v3 data when available.",
            "ICP alignment is assessed by comparing public content signals to the stated target market.",
            "Intake questionnaire answers supplement scraped data for the H (Hold/Retention) section.",
            f"Audit date: {self.date_str}  |  Data source: {self.data.get('ai_insights', {}).get('data_source', 'rule_based')}",
        ]
        for note in notes:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(note)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            r.font.color.rgb = _rgb(MGRAY)

        doc.add_paragraph()
        self._hairline(doc)
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = pf.add_run(f"C.A.S.H. Report by GMG  ·  {self.config.agency_name}  ·  {self.date_str}  ·  Confidential")
        rf.font.name = "Calibri"; rf.font.size = Pt(8)
        rf.font.color.rgb = _rgb(MGRAY)

    # ── Shared helpers ─────────────────────────────────────────

    def _section_header(self, doc: Document, title: str):
        """Navy bar with white bold text — main section divider."""
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.rows[0].cells[0]
        _shade_cell(c, NAVY)
        _cell_para(c, title, bold=True, color=WHITE, size=13,
                   align=WD_ALIGN_PARAGRAPH.LEFT)

    def _cash_section_banner(self, doc: Document, letter: str, name: str,
                              score: int, subtitle: str):
        """Two-column banner: gold letter+name left, score right."""
        sc = _score_color(score)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        lc = t.rows[0].cells[0]; rc = t.rows[0].cells[1]
        lc.width = Inches(4.5); rc.width = Inches(2.0)
        _shade_cell(lc, NAVY); _shade_cell(rc, sc)

        p1 = lc.paragraphs[0]
        p1.paragraph_format.space_before = Pt(6)
        p1.paragraph_format.space_after  = Pt(2)
        rl = p1.add_run(f"{letter}  —  {name}")
        rl.bold = True; rl.font.name = "Calibri"; rl.font.size = Pt(16)
        rl.font.color.rgb = _rgb(GOLD)
        p2 = lc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(6)
        rs = p2.add_run(subtitle)
        rs.font.name = "Calibri"; rs.font.size = Pt(9)
        rs.font.color.rgb = _rgb(MGRAY)

        _cell_para(rc, f"{score}/100\n{_grade(score)}",
                   bold=True, color=WHITE, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    def _subsection(self, doc: Document, title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(title.upper())
        r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(9)
        r.font.color.rgb = _rgb(NAVY)

    def _callout(self, doc: Document, text: str, color: str):
        """Colored left-border callout box."""
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        bar = t.rows[0].cells[0]; body = t.rows[0].cells[1]
        bar.width = Inches(0.12); body.width = Inches(6.38)
        _shade_cell(bar, color)
        bar.paragraphs[0].text = ""
        p = body.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Pt(6)
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = _rgb(DGRAY)
        doc.add_paragraph()

    def _detail_table(self, doc: Document, headers: List[str],
                      rows: List[Tuple], col_widths: List = None):
        """Standard data table: navy header row, alternating body rows."""
        n_cols = len(headers)
        t = doc.add_table(rows=1 + len(rows), cols=n_cols)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        hr = t.rows[0]
        for i, hdr in enumerate(headers):
            cell = hr.cells[i]
            if col_widths and i < len(col_widths):
                cell.width = col_widths[i]
            _shade_cell(cell, NAVY)
            _cell_para(cell, hdr, bold=True, color=WHITE, size=9,
                       align=WD_ALIGN_PARAGRAPH.LEFT)

        # Data rows
        for row_idx, row_data in enumerate(rows):
            bg = LGRAY if row_idx % 2 == 0 else WHITE
            dr = t.rows[row_idx + 1]
            for col_idx, val in enumerate(row_data):
                cell = dr.cells[col_idx]
                if col_widths and col_idx < len(col_widths):
                    cell.width = col_widths[col_idx]
                _shade_cell(cell, bg)
                text = str(val) if val is not None else "—"
                _cell_para(cell, text, bold=False, color=DGRAY, size=9)

        doc.add_paragraph()

    def _issues_strengths(self, doc: Document, issues: List[str], strengths: List[str]):
        if not issues and not strengths:
            return
        self._subsection(doc, "Issues & Strengths")
        max_rows = max(len(issues), len(strengths), 1)
        t = doc.add_table(rows=max_rows + 1, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header
        _shade_cell(t.rows[0].cells[0], RED)
        _shade_cell(t.rows[0].cells[1], GREEN)
        _cell_para(t.rows[0].cells[0], "ISSUES TO FIX", bold=True, color=WHITE, size=9)
        _cell_para(t.rows[0].cells[1], "STRENGTHS", bold=True, color=WHITE, size=9)
        t.rows[0].cells[0].width = t.rows[0].cells[1].width = Inches(3.0)

        for i in range(max_rows):
            lc = t.rows[i + 1].cells[0]
            rc = t.rows[i + 1].cells[1]
            lc.width = rc.width = Inches(3.0)
            _shade_cell(lc, LGRAY if i % 2 == 0 else WHITE)
            _shade_cell(rc, LGRAY if i % 2 == 0 else WHITE)
            _cell_para(lc, issues[i]   if i < len(issues)    else "", size=9)
            _cell_para(rc, strengths[i] if i < len(strengths) else "", size=9)

        doc.add_paragraph()

    def _hairline(self, doc: Document):
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        c = t.rows[0].cells[0]
        _shade_cell(c, MGRAY)
        c.paragraphs[0].text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

    # ── CASH score computation ─────────────────────────────────

    def _compute_cash_scores(self) -> Dict:
        ai = self.data.get("ai_insights", {})
        # Use AI-computed scores if available
        if ai.get("cash_c_score") is not None:
            return {
                "C": ai.get("cash_c_score", 50),
                "A": ai.get("cash_a_score", 50),
                "S": ai.get("cash_s_score", 50),
                "H": ai.get("cash_h_score", 50),
                "overall": ai.get("overall_score", 50),
            }

        # Compute locally
        icp   = self.data.get("icp", {}).get("score", 50)
        brand = self.data.get("brand", {}).get("score", 50)
        fresh = self.data.get("freshness", {}).get("score", 50)
        seo   = self.data.get("seo", {}).get("score", 50)
        web   = self.data.get("website", {}).get("scores", {})
        funnel_stages = self.data.get("funnel", {}).get("stages", {})
        social_n = len(self.config.active_social_channels)

        def _stage_score(stage):
            n_crit = len([i for i in stage.get("issues", [])
                          if "🔴" in i and "could not be verified" not in i
                          and "unknown" not in i.lower()])
            n_str  = len(stage.get("strengths", []))
            base   = 50 - (n_crit * 10) + (n_str * 6)
            return max(25, min(100, base)) if n_crit > 0 else max(50, min(100, 50 + n_str * 6))

        c = round((fresh + seo + web.get("content", 50)) / 3)
        a = round((icp + brand + min(social_n * 15, 80)) / 3)
        s = round((_stage_score(funnel_stages.get("capture",    {})) +
                   _stage_score(funnel_stages.get("conversion", {})) +
                   web.get("conversion", 50)) / 3)
        h = round((_stage_score(funnel_stages.get("nurture", {})) +
                   _stage_score(funnel_stages.get("trust",   {}))) / 2)

        overall = ai.get("overall_score") or round(c * 0.20 + a * 0.30 + s * 0.30 + h * 0.20)
        return {"C": c, "A": a, "S": s, "H": h, "overall": overall}
